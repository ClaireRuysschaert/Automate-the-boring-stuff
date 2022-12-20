# Document object = entire document
    # contains a list of paragraph objects 
        # a paragraph begin with "enter" or "return"
        # contains a list of Run objects

#Reading word doc
import docx
doc = docx.Document('demo.docx')
# Run objects have text attributes = text 
doc.paragraphs[1].text
#Paragraphs objects have runs attribute = list of Run objects
len(doc.paragraphs[1].runs)

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))
