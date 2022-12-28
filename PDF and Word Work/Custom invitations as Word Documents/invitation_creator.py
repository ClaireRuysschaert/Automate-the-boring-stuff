import docx

doc = docx.Document('Custom invitations.docx') 

with open('guests.txt') as file:
    for line in file:
        doc.add_paragraph('It would be a pleasure to have your company to our new eve party')
        doc.add_paragraph(line, 'Heading 1')
        p = doc.add_paragraph("at the TDB house at 19 o'clock")
        run = p.add_run()
        run.add_break(docx.enum.text.WD_BREAK.PAGE)
        

doc.save('invitations_created.docx')
