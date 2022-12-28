import docx

doc = docx.Document() #blank document
doc.add_paragraph('Hello, world!')
doc.save('helloworld.docx')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.', 'Title')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.add_heading('Header 3', 3) #from 0 to 4
doc.save('multipleParagraphs.docx')

# new Paragraph objects can be added only to the end of the document
# new Run objects can be added only to the end of a Paragraph object

#Breaks (saut de page)
doc2 = docx.Document()
doc2.add_paragraph('This is on the first page!')
p = doc2.add_paragraph(line)
run = p.add_run()
run.add_break(docx.enum.text.WD_BREAK.PAGE)
doc2.add_paragraph('This is on the second page!')
doc2.add_picture('pict.jpg', width=docx.shared.Inches(4), height=docx.shared.Cm(4))
doc2.save('twoPage.docx')
